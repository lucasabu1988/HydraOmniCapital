"""
Convierte OMNICAPITAL_MANIFESTO.md a formato DOCX
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_md_to_docx(md_file, docx_file):
    # Leer archivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Procesar línea por línea
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Título principal (nivel 1)
        if line.startswith('# '):
            text = line[2:]
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtítulo nivel 1
        elif line.startswith('## '):
            text = line[3:]
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtítulo nivel 2
        elif line.startswith('### '):
            text = line[4:]
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 76, 153)
        
        # Subtítulo nivel 3
        elif line.startswith('#### '):
            text = line[5:]
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(51, 102, 153)
        
        # Línea divisoria
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        
        # Código/bloque de código
        elif line.startswith('```'):
            # Inicio de bloque de código
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Agregar código con estilo
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.style = 'Intense Quote'
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0)
        
        # Lista numerada
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Lista con viñetas
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Checkbox
        elif line.startswith('□'):
            text = line[1:]
            p = doc.add_paragraph('☐' + text)
        
        # Tabla (formato markdown)
        elif '|' in line and not line.startswith('|---'):
            # Detectar tabla
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                if not lines[i].startswith('|---'):
                    table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 2:
                # Procesar tabla
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|') if c.strip()]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # Crear tabla
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_cells in enumerate(rows):
                        row = table.rows[row_idx]
                        for col_idx, cell_text in enumerate(row_cells):
                            if col_idx < len(row.cells):
                                row.cells[col_idx].text = cell_text
            
            continue  # Ya incrementamos i
        
        # Texto en negrita dentro de línea
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Partes entre **
                    run.bold = True
        
        # Texto en cursiva
        elif '*' in line and not line.startswith('*'):
            p = doc.add_paragraph()
            parts = re.split(r'\*(.*?)\*', line)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.italic = True
        
        # Línea normal
        elif line.strip():
            p = doc.add_paragraph(line)
        
        # Línea vacía (salto)
        else:
            pass  # Ignorar líneas vacías múltiples
        
        i += 1
    
    # Guardar documento
    doc.save(docx_file)
    print(f"Documento guardado: {docx_file}")

if __name__ == "__main__":
    convert_md_to_docx('OMNICAPITAL_MANIFESTO.md', 'OMNICAPITAL_MANIFESTO.docx')
