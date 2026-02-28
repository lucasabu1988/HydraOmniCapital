"""Export COMPASS v8 Manifesto to Word (.docx) format"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# ============================================================================
# STYLES
# ============================================================================

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()


def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ============================================================================
# COVER PAGE
# ============================================================================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('OMNICAPITAL v8')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('COMPASS')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Cross-sectional Momentum\nPosition-Adjusted Risk Scaling')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Manifiesto del Algoritmo\nVersion 8.0 | Febrero 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ============================================================================
# 1. FILOSOFIA
# ============================================================================

doc.add_heading('1. Filosofia', level=1)

doc.add_paragraph(
    'COMPASS nace de una leccion dolorosa: el v6 original reportaba 16.92% CAGR, '
    'pero al eliminar el sesgo de supervivencia, la realidad era 5.40% con -59.4% '
    'de drawdown maximo. No habia signal. Era compra aleatoria con leverage. '
    'El "alpha" era una ilusion.'
)

doc.add_paragraph(
    'COMPASS reemplaza la aleatoriedad con un edge real basado en tres pilares '
    'academicos probados en decadas de investigacion:'
)

bullets = [
    'Los ganadores siguen ganando (Momentum cross-seccional)',
    'Cuando el mercado cae, hay que salir (Filtro de regimen)',
    'La volatilidad es el verdadero riesgo (Vol targeting)',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    for run in p.runs:
        run.bold = True

doc.add_paragraph(
    'La simplicidad sigue siendo el principio rector. COMPASS usa solo datos OHLCV '
    'diarios, un indicador de mercado (SPY vs SMA200), y una formula de scoring de '
    'dos componentes. No hay machine learning, no hay optimizacion de 50 parametros, '
    'no hay caja negra.'
)

# ============================================================================
# 2. RESULTADOS
# ============================================================================

doc.add_heading('2. Resultados (Backtest 2000-2026)', level=1)

add_table(doc,
    ['Metrica', 'v6 (aleatorio)', 'v8 COMPASS', 'Mejora'],
    [
        ['CAGR', '5.40%', '16.16%', '+10.76%'],
        ['Sharpe', '0.22', '0.73', '3.3x mejor'],
        ['Sortino', '—', '1.02', '—'],
        ['Max Drawdown', '-59.4%', '-34.8%', '24.6% menos'],
        ['Calmar', '—', '0.46', '—'],
        ['Win Rate', '—', '55.25%', '—'],
        ['Trades (26 anos)', '—', '5,386', '~207/ano'],
        ['Anos positivos', '—', '21/26', '81%'],
        ['$100k se convierte en', '~$400k', '$4.95M', '12x mas'],
        ['Mejor ano', '—', '+128.1%', '—'],
        ['Peor ano', '—', '-32.7%', '—'],
    ])

# ============================================================================
# 3. EL ALGORITMO
# ============================================================================

doc.add_heading('3. El Algoritmo', level=1)

# 3.1
doc.add_heading('3.1 Universo de Inversion', level=2)

doc.add_paragraph(
    'Pool amplio: 113 acciones del S&P 500 distribuidas en 9 sectores '
    '(Technology, Financials, Healthcare, Consumer, Energy, Industrials, '
    'Utilities, Real Estate, Telecom).'
)

doc.add_paragraph(
    'Rotacion anual: Cada 1 de enero, las 113 acciones se rankean por volumen '
    'promedio diario en dolares (Close x Volume) del ano anterior. Solo las top 40 '
    'son elegibles para ese ano. Esto evita sesgo de supervivencia: no se puede '
    'invertir en acciones que solo sabemos que son grandes "hoy".'
)

doc.add_paragraph(
    'Filtro de antiguedad: Una accion necesita al menos 63 dias de historia '
    '(3 meses) para ser elegible. Esto evita la volatilidad excesiva de IPOs recientes.'
)

p = doc.add_paragraph()
run = p.add_run('Resultado historico: ')
run.bold = True
p.add_run('78 acciones unicas utilizadas en 26 anos, ~4-5 rotan cada ano.')

# 3.2
doc.add_heading('3.2 Filtro de Regimen de Mercado', level=2)

doc.add_paragraph('El sistema opera en dos modos:')

add_table(doc,
    ['Condicion', 'Regimen', 'Posiciones', 'Leverage'],
    [
        ['SPY > SMA(200) por 3+ dias', 'RISK_ON', '5', 'Vol targeting (0.5x-2.0x)'],
        ['SPY < SMA(200) por 3+ dias', 'RISK_OFF', '2', '1.0x fijo'],
    ])

doc.add_paragraph(
    'El filtro SMA200 es el indicador de tendencia mas simple y robusto que existe. '
    'Meb Faber (2007) demostro que estar fuera del mercado cuando SPY esta debajo '
    'de su SMA200 reduce el max drawdown a la mitad con minimo impacto en retornos.'
)

doc.add_paragraph(
    'Confirmacion de 3 dias: Evita whipsaw. No se cambia de regimen por un solo dia '
    'de cruce. Se requieren 3 dias consecutivos para confirmar el cambio.'
)

p = doc.add_paragraph()
run = p.add_run('Resultado historico: ')
run.bold = True
p.add_run('~25.7% del tiempo en RISK_OFF. Evito la mayoria de 2001-2002, 2008-2009, y partes de 2022.')

# 3.3
doc.add_heading('3.3 Seleccion de Acciones: El Score COMPASS', level=2)

doc.add_paragraph(
    'La pieza central del algoritmo. Cada dia, para cada accion elegible, se calcula:'
)

add_code_block(doc,
    'momentum_90d = (Precio hace 5 dias / Precio hace 90 dias) - 1\n'
    'reversal_5d  = (Precio hoy / Precio hace 5 dias) - 1\n'
    '\n'
    'SCORE = momentum_90d - reversal_5d'
)

doc.add_paragraph('Se seleccionan las acciones con mayor score (top N, donde N = posiciones disponibles).')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Que significa un score alto:')
run.bold = True

bullets = [
    ('momentum_90d alto: ', 'La accion ha subido fuerte en los ultimos 3 meses (excluyendo la ultima semana). Es un ganador de mediano plazo.'),
    ('reversal_5d bajo (o negativo): ', 'La accion tuvo un pullback reciente en la ultima semana.'),
    ('Combinacion: ', 'Ganador de mediano plazo + pullback reciente = oportunidad de compra en la tendencia.'),
]
for bold_part, normal_part in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.bold = True
    p.add_run(normal_part)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Base academica:')
run.bold = True

doc.add_paragraph(
    'Jegadeesh & Titman (1993): Momentum cross-seccional funciona en horizontes '
    'de 3-12 meses. Acciones ganadoras siguen ganando.', style='List Bullet'
)
doc.add_paragraph(
    'Lo & MacKinlay (1990): Retornos de corto plazo (1-5 dias) muestran reversion '
    'a la media. Un pullback reciente es temporal.', style='List Bullet'
)
doc.add_paragraph(
    'El "skip" de los ultimos 5 dias es critico: elimina el efecto de micro-reversion '
    'que cancela el momentum. Este es un hallazgo robusto de la literatura.', style='List Bullet'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Diferencia con v6: ')
run.bold = True
p.add_run('v6 seleccionaba al azar. COMPASS rankea y elige los mejores. '
          'Esto es la diferencia entre tirar dados y leer el mercado.')

p = doc.add_paragraph()
run = p.add_run('Diferencia con el filtro SMA20/SMA50 que fallo: ')
run.bold = True
p.add_run('Ese era un filtro de tendencia individual (trend-following por accion). '
          'COMPASS es momentum cross-seccional: compara acciones ENTRE SI y elige '
          'las mejores relativas. Son conceptos fundamentalmente diferentes.')

# 3.4
doc.add_heading('3.4 Position Sizing: Inverse Volatility', level=2)

doc.add_paragraph(
    'No todas las posiciones son iguales. Una accion con 40% de volatilidad anual '
    'no deberia tener el mismo peso que una con 15%.'
)

add_code_block(doc,
    'vol_20d(stock) = desv. estandar retornos diarios (20d) x sqrt(252)\n'
    'peso_raw(stock) = 1 / vol_20d(stock)\n'
    'peso(stock) = peso_raw / suma(todos los pesos_raw)\n'
    'tamano_posicion = peso x capital_efectivo'
)

doc.add_paragraph(
    'Efecto: Acciones estables (JNJ, PG, KO) reciben mas capital. Acciones volatiles '
    '(TSLA, NVDA, AMD) reciben menos. Esto reduce la volatilidad total del portfolio '
    'sin sacrificar retornos.'
)

doc.add_paragraph(
    'Limite: Ninguna posicion puede exceder 40% del cash disponible, independientemente de los pesos.'
)

# 3.5
doc.add_heading('3.5 Leverage Dinamico: Volatility Targeting', level=2)

doc.add_paragraph(
    'En lugar de leverage fijo (2x siempre o 1x siempre), COMPASS ajusta el leverage automaticamente:'
)

add_code_block(doc,
    'realized_vol = volatilidad realizada de SPY (20 dias) anualizada\n'
    'leverage = 15% / realized_vol\n'
    'leverage = max(0.5, min(2.0, leverage))'
)

add_table(doc,
    ['Volatilidad del mercado', 'Leverage', 'Interpretacion'],
    [
        ['8% (calma extrema)', '1.88x', 'Mercado tranquilo, apalancar'],
        ['12% (normal bajo)', '1.25x', 'Condiciones favorables'],
        ['15% (normal)', '1.00x', 'Neutro'],
        ['20% (elevada)', '0.75x', 'Cautela'],
        ['30% (crisis)', '0.50x', 'Minima exposicion'],
    ])

doc.add_paragraph(
    'La volatilidad se agrupa (volatility clustering). Dias de alta vol son seguidos '
    'por mas dias de alta vol. Reducir exposicion cuando la vol sube evita las peores '
    'perdidas. Aumentarla cuando baja captura los mejores rallies.'
)

doc.add_paragraph(
    'Solo en RISK_ON: En RISK_OFF, el leverage es siempre 1.0x. En protection mode, '
    'es 0.5x (stage 1) o 1.0x (stage 2).'
)

# 3.6
doc.add_heading('3.6 Reglas de Salida', level=2)

doc.add_paragraph('Tres mecanismos de exit, el primero que se active:')

add_table(doc,
    ['Mecanismo', 'Condicion', 'Proposito'],
    [
        ['Hold time', '>= 5 dias de trading', 'Capturar momentum, luego rotar'],
        ['Position stop', 'Retorno <= -8%', 'Limitar perdida individual'],
        ['Trailing stop', 'Subio >5%, luego cae 3% desde max', 'Proteger ganancias'],
    ])

doc.add_paragraph('Adicionalmente:')
doc.add_paragraph('Si una accion sale del top-40 anual, se cierra.', style='List Bullet')
doc.add_paragraph(
    'Si el regimen cambia a RISK_OFF, se cierran las posiciones con peor rendimiento '
    'hasta tener max 2.', style='List Bullet'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Distribucion historica de exits:')
run.bold = True

add_table(doc,
    ['Razon de salida', 'Porcentaje', 'Trades'],
    [
        ['Hold expirado', '87.0%', '4,686'],
        ['Position stop', '6.5%', '350'],
        ['Trailing stop', '5.0%', '268'],
        ['Portfolio stop', '0.9%', '46'],
        ['Regimen reduce', '0.5%', '26'],
        ['Rotacion universo', '0.2%', '10'],
    ])

# 3.7
doc.add_heading('3.7 Portfolio Stop Loss y Recovery', level=2)

p = doc.add_paragraph()
run = p.add_run('Trigger: ')
run.bold = True
p.add_run('Drawdown del portfolio >= -15% desde el peak.')

p = doc.add_paragraph()
run = p.add_run('Accion inmediata: ')
run.bold = True
p.add_run('Cerrar TODAS las posiciones. Entrar en modo proteccion.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Recovery gradual en 2 etapas:')
run.bold = True

add_table(doc,
    ['Etapa', 'Condicion', 'Max Posiciones', 'Leverage'],
    [
        ['Stage 1', 'Primeros 63 dias post-stop', '2', '0.5x'],
        ['Stage 2', '63 dias + regimen RISK_ON', '3', '1.0x'],
        ['Normal', '126 dias + regimen RISK_ON', '5', 'Vol targeting'],
    ])

doc.add_paragraph(
    'Requisito critico: Cada etapa de recovery requiere que el mercado este en '
    'RISK_ON (SPY > SMA200). Si el mercado sigue en bear, el sistema NO restaura '
    'leverage aunque haya pasado el tiempo.'
)

# ============================================================================
# 4. COSTOS
# ============================================================================

doc.add_heading('4. Costos y Friccion', level=1)

add_table(doc,
    ['Concepto', 'Costo', 'Notas'],
    [
        ['Margin', '6% anual sobre borrowed', 'Solo cuando leverage > 1.0x'],
        ['Commission', '$0.001 por accion', 'Realista para IBKR'],
        ['Hedge cost', 'ELIMINADO', 'Vol targeting actua como hedge natural'],
        ['Slippage', 'No modelado', 'Mitigado por stocks liquidos (top-40)'],
    ])

# ============================================================================
# 5. PARAMETROS
# ============================================================================

doc.add_heading('5. Parametros Completos', level=1)

add_code_block(doc,
    'UNIVERSO\n'
    '  broad_pool          = 113 stocks (S&P 500 multi-sector)\n'
    '  top_n               = 40 (rotacion anual por dollar volume)\n'
    '  min_age_days        = 63 (3 meses minimo de historia)\n'
    '\n'
    'SIGNAL\n'
    '  momentum_lookback   = 90 dias\n'
    '  momentum_skip       = 5 dias\n'
    '\n'
    'REGIMEN\n'
    '  regime_sma_period   = 200 dias (SPY)\n'
    '  regime_confirm_days = 3\n'
    '\n'
    'POSICIONES\n'
    '  num_positions       = 5 (RISK_ON)\n'
    '  num_positions_off   = 2 (RISK_OFF)\n'
    '  hold_days           = 5 (dias de trading)\n'
    '\n'
    'RIESGO POR POSICION\n'
    '  position_stop_loss  = -8%\n'
    '  trailing_activation = +5%\n'
    '  trailing_stop_pct   = -3%\n'
    '\n'
    'RIESGO DE PORTFOLIO\n'
    '  portfolio_stop_loss = -15%\n'
    '  recovery_stage_1    = 63 dias + RISK_ON\n'
    '  recovery_stage_2    = 126 dias + RISK_ON\n'
    '\n'
    'LEVERAGE\n'
    '  target_vol          = 15% anualizado\n'
    '  leverage_min        = 0.5x\n'
    '  leverage_max        = 2.0x\n'
    '  vol_lookback        = 20 dias\n'
    '\n'
    'COSTOS\n'
    '  initial_capital     = $100,000\n'
    '  margin_rate         = 6% anual\n'
    '  commission          = $0.001/accion'
)

# ============================================================================
# 6. FLUJO DIARIO
# ============================================================================

doc.add_heading('6. Flujo Diario de Operacion', level=1)

steps = [
    ('VALORAR', 'Calcular valor del portfolio (cash + posiciones a mercado). Actualizar peak si hay nuevo maximo.'),
    ('PROTECCION', 'Si en recovery: verificar condiciones de siguiente etapa. Calcular drawdown. Si drawdown <= -15%: STOP LOSS, cerrar todo, entrar Stage 1.'),
    ('REGIMEN', 'Leer precio de SPY. Comparar con SMA(200). Si 3+ dias consecutivos en nuevo lado: cambiar regimen.'),
    ('CERRAR POSICIONES', 'En orden: (a) Hold >= 5 dias, (b) Retorno <= -8%, (c) Trailing activado y caida >= 3%, (d) Fuera del top-40 anual, (e) Exceso de posiciones por cambio de regimen.'),
    ('ABRIR POSICIONES', 'Calcular SCORE para cada accion elegible. Rankear. Seleccionar top N. Calcular pesos inverse-vol. Calcular leverage por vol targeting. Abrir posiciones.'),
    ('COSTOS', 'Deducir margin cost diario si leverage > 1.0x.'),
    ('REGISTRAR', 'Snapshot diario: valor, cash, posiciones, drawdown, leverage, regimen.'),
]

for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================================
# 7. ESCENARIOS HISTORICOS
# ============================================================================

doc.add_heading('7. Comportamiento en Escenarios Historicos', level=1)

scenarios = [
    ('Dot-com crash (2000-2002)',
     'Stop loss activado Sep 2000 (DD -15.5%). Recovery en Ene 2002. '
     'Segundo stop en Abr 2002. Recovery en Abr 2003. '
     'Proteccion evito lo peor del crash.'),
    ('Bull market 2003-2007',
     'Crecimiento sostenido con leverage 1.5x-2.0x. '
     'Stops menores en 2004 y 2006, recovery rapida en ~6 meses cada vez.'),
    ('Crisis financiera 2008-2009',
     'Stop loss en Feb 2008 (antes del crash principal). '
     'RISK_OFF activo la mayor parte de 2008. Recovery en Jun 2009. '
     'Evito la caida de -50%+ del mercado general.'),
    ('Bull market 2010-2019',
     '10 anos de crecimiento. Stops menores en 2011 y 2012. '
     'Leverage entre 1.5x-2.0x la mayoria del tiempo. '
     'Portfolio crece de ~$220k a $1M+.'),
    ('COVID crash Mar 2020',
     'Stop loss 9 Mar 2020 (DD -17.7%). Recovery Stage 1 Jun 2020. '
     'Stage 2 Sep 2020. Rally post-COVID captura retornos excepcionales. '
     'Mejor ano: +128%.'),
    ('Bear market 2022',
     'Stop loss Ene 2022 (DD -15%). RISK_OFF activo gran parte del ano. '
     'Recovery completa en May 2023.'),
    ('2024-2026',
     'Stop en Jul 2024, recovery en Ene 2025. Stop en Mar 2025, recovery en Ene 2026. '
     'Valor final: ~$4.95M.'),
]

for title, desc in scenarios:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ============================================================================
# 8. RIESGOS
# ============================================================================

doc.add_heading('8. Riesgos y Limitaciones', level=1)

doc.add_heading('Riesgos conocidos', level=2)

risks = [
    ('Overfitting', 'Los parametros fueron elegidos con base academica, no optimizados sobre este dataset. Sin embargo, cualquier backtest tiene riesgo de overfitting. Se necesita out-of-sample testing.'),
    ('Momentum crash', 'El momentum como factor puede sufrir "crashes" rapidos (tipicamente en recuperaciones de mercado como Mar 2009). El regime filter mitiga esto parcialmente.'),
    ('Costos reales', 'El slippage no esta modelado explicitamente. En acciones liquidas del top-40 deberia ser minimo, pero en dias de alta volatilidad podria ser significativo.'),
    ('Dependencia de un indicador', 'El filtro de regimen depende de SPY vs SMA200. Si el mercado cambia de estructura, el filtro no protege.'),
    ('Max drawdown de -34.8%', 'Aunque mejor que v6 (-59.4%), sigue siendo significativo. Un inversor conservador podria no tolerar una caida de un tercio.'),
]

for title, desc in risks:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Lo que NO hace COMPASS', level=2)

nots = [
    'No predice el futuro',
    'No usa machine learning ni AI',
    'No opera intradayia',
    'No hace short selling',
    'No usa opciones ni derivados',
    'No analiza fundamentales ni noticias',
    'No promete retornos garantizados',
]
for n in nots:
    doc.add_paragraph(n, style='List Bullet')

# ============================================================================
# 9. IMPLEMENTACION
# ============================================================================

doc.add_heading('9. Implementacion Tecnica', level=1)

p = doc.add_paragraph()
run = p.add_run('Archivo principal: ')
run.bold = True
p.add_run('omnicapital_v8_compass.py (~875 lineas de Python)')

p = doc.add_paragraph()
run = p.add_run('Dependencias: ')
run.bold = True
p.add_run('pandas, numpy (calculo), yfinance (datos), pickle (cache)')

p = doc.add_paragraph()
run = p.add_run('Para ejecutar: ')
run.bold = True

add_code_block(doc, 'python omnicapital_v8_compass.py')

# ============================================================================
# 10. REGLAS INQUEBRANTABLES
# ============================================================================

doc.add_heading('10. Reglas Inquebrantables', level=1)

rules = [
    'NO modificar los parametros sin backtest completo. Cada cambio debe probarse en todo el periodo 2000-2026.',
    'NO agregar complejidad. Si una mejora no aporta al menos +1% CAGR o -5% max DD, no vale la pena.',
    'NO ignorar el regime filter. La tentacion de operar en RISK_OFF "porque esta vez es diferente" es la causa #1 de perdidas catastroficas.',
    'NO aumentar leverage max por encima de 2.0x. El vol targeting ya optimiza el leverage.',
    'NO operar acciones fuera del top-40. El universo esta definido por liquidez.',
    'SIEMPRE respetar los stops. Position stop (-8%), trailing stop (-3% desde max), portfolio stop (-15%). Sin excepciones.',
    'Paper trading primero. Minimo 3 meses de paper trading antes de capital real.',
]

for i, rule in enumerate(rules, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    # Split bold part and normal part
    parts = rule.split('. ', 1)
    if len(parts) == 2:
        run2 = p.add_run(parts[0] + '. ')
        run2.bold = True
        p.add_run(parts[1])
    else:
        p.add_run(rule)

# ============================================================================
# 11. EVOLUCION
# ============================================================================

doc.add_heading('11. Evolucion desde v1', level=1)

add_table(doc,
    ['Version', 'Fecha', 'Descripcion', 'CAGR', 'Sharpe'],
    [
        ['v1', 'Feb 2026', 'MicroManagement intradayia', '—', '—'],
        ['v4', 'Feb 2026', 'Optimized daily', '—', '—'],
        ['v5', 'Feb 2026', '3 Day Strategy', '—', '—'],
        ['v6', 'Feb 2026', 'Random selection, 2x leverage', '5.40%*', '0.22'],
        ['v8 COMPASS', 'Feb 2026', 'Momentum + Regime + Vol targeting', '16.16%', '0.73'],
    ])

doc.add_paragraph(
    '*v6 CAGR con universo corregido (sin sesgo de supervivencia). '
    'El valor original de 16.92% incluia sesgo.'
)

# ============================================================================
# 12. CONCLUSION
# ============================================================================

doc.add_heading('12. Conclusion', level=1)

doc.add_paragraph(
    'COMPASS demuestra que reemplazar la aleatoriedad con un signal academicamente '
    'fundamentado, combinado con gestion de riesgo adaptativa, puede triplicar los '
    'retornos ajustados por riesgo.'
)

doc.add_paragraph(
    'El algoritmo no es perfecto. Tiene 11 stop loss events en 26 anos y un max '
    'drawdown de -34.8%. Pero a diferencia del v6, cada componente tiene una razon de ser:'
)

components = [
    ('Momentum', 'Compra ganadores → edge positivo'),
    ('Regime filter', 'Sale en bear markets → protege capital'),
    ('Vol targeting', 'Adapta exposicion → suaviza la curva'),
    ('Position stops', 'Corta perdedores rapido → limita dano'),
    ('Inverse vol sizing', 'Equilibra riesgo → reduce volatilidad'),
]
for name, desc in components:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'Cinco mecanismos simples, cada uno con decadas de evidencia academica. '
    'Juntos, transforman $100k en $4.95M en 26 anos.'
)

doc.add_paragraph()
quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = quote.add_run('"Don\'t confuse randomness with edge. COMPASS knows the difference."')
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

# ============================================================================
# SAVE
# ============================================================================

output_path = 'OMNICAPITAL_V8_COMPASS_MANIFESTO.docx'
doc.save(output_path)
print(f"\nManifiesto exportado: {output_path}")
